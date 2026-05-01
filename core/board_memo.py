"""
JCL Debt Monitor — Board Memo Generator
Creates a downloadable Word document summarizing portfolio status for board / management review.
"""

from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ui.insights import (
    calculate_health_score, generate_recommendations, generate_bottom_line,
)


# Color palette (RGB)
NAVY    = RGBColor(0x0F, 0x1E, 0x35)
SLATE   = RGBColor(0x1E, 0x29, 0x3B)
BLUE    = RGBColor(0x3B, 0x82, 0xF6)
GREEN   = RGBColor(0x10, 0xB9, 0x81)
AMBER   = RGBColor(0xF5, 0x9E, 0x0B)
RED     = RGBColor(0xEF, 0x44, 0x44)
GRAY    = RGBColor(0x64, 0x74, 0x8B)
TEXT    = RGBColor(0x1F, 0x29, 0x37)


def _set_cell_background(cell, color_hex: str):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_heading(doc, text: str, level: int = 1, color: RGBColor = NAVY):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_body(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = TEXT
    return p


def _add_kpi_table(doc, kpis: list):
    """Add a 4-column KPI table."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row (labels)
    for i, kpi in enumerate(kpis):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(kpi["label"].upper())
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_background(cell, "F1F5F9")

    # Value row
    for i, kpi in enumerate(kpis):
        cell = table.cell(1, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(kpi["value"])
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = kpi.get("color", NAVY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "sub" in kpi and kpi["sub"]:
            sub_p = cell.add_paragraph()
            sub_run = sub_p.add_run(kpi["sub"])
            sub_run.font.name = "Calibri"
            sub_run.font.size = Pt(8)
            sub_run.font.color.rgb = GRAY
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_background(cell, "FFFFFF")

    # Set column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.6)


def _add_recommendation(doc, rec: dict):
    """Add a single recommendation as a styled box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""

    color_map = {
        "HIGH":   ("FEE2E2", "991B1B"),
        "MEDIUM": ("FEF3C7", "92400E"),
        "LOW":    ("DBEAFE", "1E40AF"),
        "INFO":   ("F1F5F9", "475569"),
    }
    bg, _ = color_map.get(rec["priority"], ("F1F5F9", "475569"))
    _set_cell_background(cell, bg)

    # Priority + Title
    p = cell.paragraphs[0]
    pr = p.add_run(f"[{rec['priority']}] ")
    pr.bold = True
    pr.font.name = "Calibri"
    pr.font.size = Pt(9)
    if rec["priority"] == "HIGH": pr.font.color.rgb = RED
    elif rec["priority"] == "MEDIUM": pr.font.color.rgb = AMBER
    elif rec["priority"] == "LOW": pr.font.color.rgb = BLUE
    else: pr.font.color.rgb = GRAY

    tr = p.add_run(rec["title"])
    tr.bold = True
    tr.font.name = "Calibri"
    tr.font.size = Pt(11)
    tr.font.color.rgb = NAVY

    # Body
    body_p = cell.add_paragraph()
    body_run = body_p.add_run(rec["body"])
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = TEXT

    # Owner
    owner_p = cell.add_paragraph()
    owner_run = owner_p.add_run(f"Owner: {rec['owner']}")
    owner_run.font.name = "Calibri"
    owner_run.font.size = Pt(8)
    owner_run.italic = True
    owner_run.font.color.rgb = GRAY

    # Spacing after the box
    doc.add_paragraph()


def _add_data_table(doc, headers: list, rows: list):
    """Add a clean data table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        _set_cell_background(cell, "DBEAFE")

    # Data
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT


def _strip_html(text: str) -> str:
    """Remove HTML tags and entities from a string."""
    import re
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&nbsp;", " ").replace("&amp;", "&")
    return text


def generate_board_memo(logic, controls) -> bytes:
    """Generate the board memo as a Word document. Returns bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === HEADER ===
    title = doc.add_paragraph()
    title_run = title.add_run("JINDAL COKE LIMITED")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Debt Portfolio — Board Memorandum")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = BLUE
    subtitle_run.italic = True

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Prepared: {date.today().strftime('%d %B %Y')}  |  "
        f"As-of: {logic.as_of_date.strftime('%d %B %Y')}  |  "
        f"Basis: {logic.basis}"
    )
    meta_run.font.name = "Calibri"
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GRAY

    # Divider
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "3B82F6")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === EXECUTIVE SUMMARY ===
    bl = generate_bottom_line(logic, controls)
    _add_heading(doc, "Executive Summary", level=1)

    verdict_color_map = {
        "HEALTHY": GREEN, "ALL CLEAR": GREEN,
        "MONITOR CLOSELY": AMBER, "MONITOR": AMBER, "RENEWALS DUE": AMBER,
        "ACTION REQUIRED": RED, "BREACH": RED,
        "ON TRACK": GREEN, "HIGHLY RESILIENT": GREEN, "RESILIENT TO MODERATE": BLUE,
    }
    verdict_color = verdict_color_map.get(bl["verdict"], BLUE)

    verdict_p = doc.add_paragraph()
    verdict_run = verdict_p.add_run(f"STATUS: {bl['verdict']}")
    verdict_run.bold = True
    verdict_run.font.name = "Calibri"
    verdict_run.font.size = Pt(12)
    verdict_run.font.color.rgb = verdict_color

    _add_body(doc, _strip_html(bl["narrative"]), size=11)

    # === KEY METRICS ===
    _add_heading(doc, "Portfolio Metrics", level=2)
    ls = logic.lender_summary()
    total_sanc = ls["Total_Sanction"].sum()
    wac = logic.weighted_avg_cost(controls["rate_shock"], controls["spread_shock"])
    annual_int = logic.calculate_annual_interest(controls["rate_shock"], controls["spread_shock"])
    cov_df = logic.calculate_covenants(controls["ebitda_change"])
    compliant = (cov_df["Status"] == "Compliant").sum()
    near = (cov_df["Status"] == "Near Breach").sum()
    breach = (cov_df["Status"] == "Breach").sum()
    health = calculate_health_score(logic, controls)
    health_color = {"#10B981": GREEN, "#3B82F6": BLUE, "#F59E0B": AMBER, "#EF4444": RED}.get(health["color"], NAVY)

    _add_kpi_table(doc, [
        {"label": "Total Portfolio", "value": f"₹{total_sanc:,.0f} Cr",
         "sub": f"{len(logic.facility_master)} facilities · 5 lenders"},
        {"label": "Annual Cost", "value": f"₹{annual_int['total']:.0f} Cr",
         "sub": f"WAC {wac*100:.2f}%"},
        {"label": "Covenants", "value": f"{compliant}/{len(cov_df)}",
         "sub": "All compliant" if (breach + near) == 0 else f"{breach} breach, {near} near", "color": health_color},
        {"label": "Health Score", "value": f"{health['total']}/100",
         "sub": health["rating"], "color": health_color},
    ])

    # === RECOMMENDED ACTIONS ===
    _add_heading(doc, "Recommended Actions", level=1)
    _add_body(doc, "The following items are auto-derived from current portfolio state and ranked by priority.",
               italic=True, size=10)

    recs = generate_recommendations(logic, controls)
    priority_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFO": 3}
    recs_sorted = sorted(recs, key=lambda r: priority_order.get(r["priority"], 99))

    for rec in recs_sorted:
        _add_recommendation(doc, rec)

    # Page break
    doc.add_page_break()

    # === LENDER BREAKDOWN ===
    _add_heading(doc, "Lender Exposure Breakdown", level=1)
    lender_rows = []
    for _, row in ls.iterrows():
        lender_rows.append([
            row["Lender"],
            f"₹{row['Total_Sanction']:,.1f}",
            f"₹{row['Outstanding']:,.1f}",
            int(row["Num_Facilities"]),
            f"{row['Weighted_Avg_Cost']*100:.2f}%",
        ])
    _add_data_table(
        doc,
        headers=["Lender", "Sanctioned (Cr)", "Outstanding (Cr)", "# Facilities", "WAC"],
        rows=lender_rows,
    )

    # === COVENANTS REQUIRING ATTENTION ===
    watch_items = cov_df[cov_df["Status"].isin(["Breach", "Near Breach", "Watch"])]
    if len(watch_items) > 0:
        _add_heading(doc, "Covenants Requiring Attention", level=1)
        cov_rows = []
        for _, row in watch_items.iterrows():
            actual_str = f"{row['Actual']:.2f}x" if isinstance(row["Actual"], (int, float)) else str(row["Actual"])[:25]
            threshold_str = f"{row['Operator']}{row['Threshold']:.2f}x" if row["Type"] != "rating" else "≥ A-"
            hr_str = f"{row['Headroom_Pct']:+.1f}%" if row["Headroom_Pct"] is not None else "—"
            cov_rows.append([
                row["Lender"], row["Covenant"], threshold_str, actual_str, hr_str, row["Status"],
            ])
        _add_data_table(
            doc,
            headers=["Lender", "Covenant", "Threshold", "Actual", "Headroom %", "Status"],
            rows=cov_rows,
        )
    else:
        _add_heading(doc, "Covenant Status", level=1)
        _add_body(doc, "All 24 covenants are compliant with healthy headroom. No watch items.", size=11)

    # === STRESS TESTING ===
    _add_heading(doc, "Stress Test Results", level=1)
    _add_body(doc, "How key ratios respond to a range of stress scenarios:", size=11)

    scenarios_to_test = [
        ("Base Case", 0, 0, 0),
        ("Rate +100 bps", 100, 0, 0),
        ("EBITDA −20%", 0, 0, -20),
        ("Combined Stress", 100, 50, -20),
        ("Severe Stress", 200, 100, -30),
    ]
    stress_rows = []
    for name, rs, ss, es in scenarios_to_test:
        sc = logic.run_scenario(rs, ss, es)
        cv = sc["stress"]["covenants"]
        dscr = cv[cv["Covenant"] == "DSCR"]["Actual"].iloc[0]
        icr = cv[cv["Covenant"] == "ICR"]["Actual"].iloc[0]
        stress_breach = (cv["Status"] == "Breach").sum()
        stress_rows.append([
            name,
            f"₹{sc['stress']['annual_interest']:.1f} Cr",
            f"{sc['stress']['weighted_avg_cost']*100:.2f}%",
            f"{dscr:.2f}x",
            f"{icr:.2f}x",
            "PASS" if stress_breach == 0 else f"{stress_breach} BREACH",
        ])
    _add_data_table(
        doc,
        headers=["Scenario", "Annual Interest", "WAC", "DSCR", "ICR", "Result"],
        rows=stress_rows,
    )

    # === FOOTER ===
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"\n— END OF MEMO —\n\n"
        f"Source: JCL Debt Model (verified against sanction letters)  ·  "
        f"Generated by JCL Debt Monitor v1.0"
    )
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_run.font.color.rgb = GRAY
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def generate_email_summary(logic, controls) -> str:
    """Generate a copy-paste ready email status summary."""
    bl = generate_bottom_line(logic, controls)
    ls = logic.lender_summary()
    total_sanc = ls["Total_Sanction"].sum()
    wac = logic.weighted_avg_cost(controls["rate_shock"], controls["spread_shock"])
    annual_int = logic.calculate_annual_interest(controls["rate_shock"], controls["spread_shock"])
    cov_df = logic.calculate_covenants(controls["ebitda_change"])
    compliant = (cov_df["Status"] == "Compliant").sum()
    breach = (cov_df["Status"] == "Breach").sum()
    near = (cov_df["Status"] == "Near Breach").sum()
    health = calculate_health_score(logic, controls)

    recs = generate_recommendations(logic, controls)
    priority_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFO": 3}
    recs_sorted = sorted(recs, key=lambda r: priority_order.get(r["priority"], 99))

    body = f"""Subject: JCL Debt Portfolio Update — {date.today().strftime('%d %b %Y')}

Hi all,

Here's a quick snapshot of our debt portfolio as of {logic.as_of_date.strftime('%d %b %Y')}:

STATUS: {bl['verdict']}

{_strip_html(bl['narrative'])}

KEY METRICS
• Total sanctioned: ₹{total_sanc:,.0f} Cr across {len(logic.facility_master)} facilities, 5 lenders
• Annual debt service cost: ₹{annual_int['total']:.0f} Cr (WAC {wac*100:.2f}%)
• Covenant compliance: {compliant}/{len(cov_df)} compliant"""

    if breach > 0 or near > 0:
        body += f" ({breach} breach, {near} near breach)"

    body += f"""
• Portfolio health score: {health['total']}/100 ({health['rating']})

PRIORITY ACTIONS
"""

    high_med = [r for r in recs_sorted if r["priority"] in ("HIGH", "MEDIUM")]
    if not high_med:
        body += "• No urgent actions — continue regular quarterly review cycle.\n"
    else:
        for rec in high_med[:5]:
            body += f"• [{rec['priority']}] {rec['title']} (Owner: {rec['owner']})\n"

    body += f"""
Full dashboard with stress testing available on request.

Best regards
"""
    return body
